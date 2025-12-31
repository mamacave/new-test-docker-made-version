from io import BytesIO
from decimal import Decimal
from typing import Dict, Any
from docx import Document
from docx.shared import Inches
import base64


def _maybe_add_logo(document: Document, logo_data_url: str):
    if not logo_data_url:
        return
    if logo_data_url.startswith("data:"):
        header, b64 = logo_data_url.split(",", 1)
        img_data = base64.b64decode(b64)
        bio = BytesIO(img_data)
        try:
            document.add_picture(bio, width=Inches(1.5))
        except Exception:
            pass


def proposal_to_docx_bytes(proposal: Dict[str, Any]) -> bytes:
    doc = Document()
    meta = proposal.get("meta", {})
    title = meta.get("title", proposal.get("name", "Proposal"))
    doc.add_heading(title, level=1)

    # logo
    _maybe_add_logo(doc, meta.get("logo_data_url"))

    pid = meta.get("proposal_id", proposal.get("proposal_id", ""))
    date = meta.get("date", proposal.get("date", ""))
    if pid or date:
        p = doc.add_paragraph()
        if pid:
            p.add_run(f"Proposal ID: {pid}").bold = True
            p.add_run("    ")
        if date:
            p.add_run(f"Date: {date}")

    doc.add_paragraph(proposal.get("notes", ""))

    addons = []
    for sec in proposal.get("sections", []):
        doc.add_heading(sec.get("title", "Section"), level=2)
        for li in sec.get("line_items", []):
            code = li.get("code")
            qty = li.get("quantity", 1)
            desc = li.get("description", code)
            doc.add_paragraph(f"- {desc} (x{qty})")
        for ao in sec.get("add_ons", []):
            code = ao.get("code")
            qty = ao.get("quantity", 1)
            doc.add_paragraph(f"* Add-on {code} (x{qty})")

    totals = proposal.get("totals")
    if not totals:
        # leave totals to caller; otherwise compute simple sum if provided
        totals = {}
    doc.add_paragraph("")
    doc.add_heading("Totals", level=2)
    for k in ("subtotal", "tax", "total"):
        v = totals.get(k)
        if v is not None:
            doc.add_paragraph(f"{k.capitalize()}: ${Decimal(v):.2f}")

    bio = BytesIO()
    doc.save(bio)
    return bio.getvalue()
